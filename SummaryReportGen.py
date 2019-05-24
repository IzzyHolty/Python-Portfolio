#Using python-docx and tkinter

import docx
from tkinter import *


#create document
reportdoc = docx.Document()

#create base window
root = Tk()
#root.geometry(geometry=500x500+200+50)
root.title("Summary Report Generator")

heading = Label(root, text="Create a Summary Report", font=("Times New Roman", 20, "bold"))
heading.pack(side=TOP)


#Form fields
requestName = str
entryName = Entry(root, textvariable=requestName, width=30, bg="lightgray")
entryName.pack(side=LEFT)


firstParagraph = reportdoc.add_paragraph("Cybersecurity recieved the request from to assess")

productName = str
entryProduct = Entry(root, text="Enter the product name.", textvariable=productName, width=30, bg="lightgray")
entryProduct.pack(side=LEFT)

firstParagraph.add_run("s security controls.")

#eventually

def submitButtonFunction():
	reportdoc.save("Summary Report.docx")

submitButton = Button(root, text="Submit", command=submitButtonFunction())
submitButton.pack(side=LEFT)

root.mainloop()