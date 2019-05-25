#Using main modules python-docx and tkinter

import docx
import random
import sys
import os
import matplotlib
import collections
import string
from tkinter import *


#text above entry box
def textEntry(words, xpos, ypos):
    textlabel = Label(root, text=words, font=("Times New Roman", 11))
    textlabel.place(x=xpos, y=ypos)

#entry field
def entryBox(xpos, ypos):
    entryName = Entry(root, width=20, bg="lightgray")
    entryName.place(x=xpos, y=ypos)
    userinput = entryName.get()
    return userinput


#create document
reportdoc = docx.Document()

#create base window
root = Tk()
root.geometry("500x500")
root.title("Document Generator")

#in-window title
heading = Label(root, text="Create a Document Template", font=("Times New Roman", 20, "bold"))
heading.pack(side=TOP)

textEntry("The person's name is", 35, 60)
requestName = entryBox(35, 80)

#write name with text concatenated to doc
firstParagraph = reportdoc.add_paragraph("This person's name is " + requestName + ".")

#text above entry box
textlabel = Label(root, text="They like", font=("Times New Roman", 11))
textlabel.place(x=35, y=120)

#enter person's thing they like
thingLike = Entry(root, width=20, bg="lightgray")
thingLike.place(x=35, y=140)
personLikes = thingLike.get()

#continue on line
firstParagraph.add_run("They like " + personLikes + ".")

#Name the document
docName = Entry(root, width=20, bg="lightgray")
docName.place(x=35, y=160)
documentname = docName.get()

#submit button
def submitButtonFunction():
	reportdoc.save(documentname + ".docx")

submitButton = Button(root, text="Submit", command=submitButtonFunction())
submitButton.place(x=35, y=180)

#main
root.mainloop()
